"""
Document Parser — Multi-format academic document parsing
Supports PDF (PyMuPDF), DOCX (python-docx), Markdown, and plain text.
Auto-detects Chinese/English and extracts structured metadata.
"""
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class ParsedDocument:
    """Parsed document structure."""
    title: str = ""
    authors: str = ""
    year: Optional[int] = None
    abstract: str = ""
    keywords: list = field(default_factory=list)
    sections: list = field(default_factory=list)  # [(section_name, content)]
    full_text: str = ""
    page_count: int = 0
    file_path: str = ""
    language: str = "zh"  # zh or en
    metadata: dict = field(default_factory=dict)


class PDFParser:
    """Academic document parser with CN/EN support."""

    # Chinese section heading patterns
    CN_SECTION_PATTERNS = [
        r'^[一二三四五六七八九十]+[、．.]',
        r'^第[一二三四五六七八九十]+[部章节]',
        r'^[（\(][一二三四五六七八九十]+[）\)]',
        r'^\d+[、．.\s]',
        r'^\d+\.\d+\s',
    ]

    # English section heading patterns
    EN_SECTION_PATTERNS = [
        r'^(?:I{1,3}|IV|V|VI{0,3})\.\s',
        r'^\d+\.\s+[A-Z]',
        r'^(?:Abstract|Introduction|Literature|Method|Result|Conclusion|Discussion|Reference)',
    ]

    def __init__(self):
        self.fitz = None
        self.docx_module = None
        try:
            import fitz  # PyMuPDF
            self.fitz = fitz
        except ImportError:
            pass
        try:
            import docx  # python-docx
            self.docx_module = docx
        except ImportError:
            pass

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document file (PDF/DOCX/MD/TXT)."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix == '.pdf':
            return self._parse_pdf(path)
        elif suffix == '.docx':
            return self._parse_docx(path)
        elif suffix in ('.md', '.txt'):
            return self._parse_text(path)
        elif suffix == '.doc':
            raise ValueError(f"Legacy .doc format not supported, convert to .docx: {file_path}")
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Parse PDF file."""
        if not self.fitz:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

        doc = self.fitz.open(str(path))
        parsed = ParsedDocument(file_path=str(path), page_count=len(doc))

        full_text_parts = []
        for page in doc:
            text = page.get_text("text")
            full_text_parts.append(text)
        parsed.full_text = "\n".join(full_text_parts)

        parsed.language = self._detect_language(parsed.full_text)

        meta = doc.metadata
        if meta:
            parsed.title = meta.get("title", "") or ""
            parsed.authors = meta.get("author", "") or ""
            parsed.metadata = {k: v for k, v in meta.items() if v}

        if not parsed.title:
            parsed.title = self._extract_title(parsed.full_text)

        parsed.abstract = self._extract_abstract(parsed.full_text, parsed.language)
        parsed.keywords = self._extract_keywords(parsed.full_text, parsed.language)
        parsed.year = self._extract_year(parsed.full_text, str(path))
        parsed.sections = self._extract_sections(parsed.full_text, parsed.language)

        doc.close()
        return parsed

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parse DOCX file (supports table-based reading notes)."""
        if not self.docx_module:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = self.docx_module.Document(str(path))
        parsed = ParsedDocument(file_path=str(path))

        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        metadata_from_table = {}
        for table in doc.tables:
            for row in table.rows:
                cell_text = row.cells[0].text.strip()
                if not cell_text:
                    continue
                if len(row.cells) >= 3:
                    key = row.cells[0].text.strip()
                    val = row.cells[1].text.strip()
                    if key in ('书名', '作者', '出版社', '引用格式', '阅读日期',
                               'Title', 'Author', 'Publisher', 'Citation', 'Date'):
                        metadata_from_table[key] = val
                        continue
                if cell_text in ('读书笔记', '总结收获', '笔记/书摘',
                                 'Reading Notes', 'Summary', 'Notes'):
                    continue
                table_texts.append(cell_text)

        all_texts = para_texts + table_texts
        parsed.full_text = "\n\n".join(all_texts)
        parsed.page_count = max(1, len(parsed.full_text) // 3000)

        parsed.language = self._detect_language(parsed.full_text)

        title_keys = ('书名', 'Title')
        author_keys = ('作者', 'Author')
        citation_keys = ('引用格式', 'Citation')

        for key in title_keys:
            if key in metadata_from_table:
                parsed.title = metadata_from_table[key]
                break
        else:
            parsed.title = path.stem

        for key in author_keys:
            if key in metadata_from_table:
                parsed.authors = metadata_from_table[key]
                break

        for key in citation_keys:
            if key in metadata_from_table:
                parsed.metadata['citation'] = metadata_from_table[key]
                break

        parsed.abstract = self._extract_abstract(parsed.full_text, parsed.language)
        parsed.keywords = self._extract_keywords(parsed.full_text, parsed.language)
        parsed.year = self._extract_year(parsed.full_text, str(path))
        parsed.sections = self._extract_sections(parsed.full_text, parsed.language)

        return parsed

    def _parse_text(self, path: Path) -> ParsedDocument:
        """Parse MD/TXT file."""
        text = path.read_text(encoding='utf-8')
        parsed = ParsedDocument(
            file_path=str(path),
            full_text=text,
            page_count=max(1, len(text) // 3000),
        )

        parsed.language = self._detect_language(text)
        parsed.title = path.stem
        parsed.abstract = self._extract_abstract(text, parsed.language)
        parsed.keywords = self._extract_keywords(text, parsed.language)
        parsed.year = self._extract_year(text, str(path))
        parsed.sections = self._extract_sections(text, parsed.language)

        return parsed

    def _detect_language(self, text: str) -> str:
        """Detect document language (zh/en)."""
        sample = text[:2000]
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', sample))
        total_chars = len(sample)
        if total_chars == 0:
            return "zh"
        return "zh" if chinese_chars / total_chars > 0.1 else "en"

    def _extract_title(self, text: str) -> str:
        """Extract title from text beginning."""
        lines = text.strip().split('\n')
        title_lines = []
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 5 and len(line) < 200:
                title_lines.append(line)
                if len(title_lines) >= 2:
                    break
        return title_lines[0] if title_lines else ""

    def _extract_abstract(self, text: str, language: str) -> str:
        """Extract abstract."""
        if language == "zh":
            patterns = [
                r'摘\s*要[：:]\s*(.*?)(?=关键词|关\s*键\s*词|Abstract)',
                r'【摘要】\s*(.*?)(?=【关键词】)',
            ]
        else:
            patterns = [
                r'Abstract[:\s]*(.*?)(?=Keywords|JEL|Introduction|\d+\.\s)',
                r'ABSTRACT[:\s]*(.*?)(?=KEYWORDS|JEL|INTRODUCTION)',
            ]

        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                abstract = match.group(1).strip()
                abstract = re.sub(r'\s+', ' ', abstract)
                return abstract[:2000]
        return ""

    def _extract_keywords(self, text: str, language: str) -> list:
        """Extract keywords."""
        if language == "zh":
            patterns = [
                r'关\s*键\s*词[：:]\s*(.*?)(?=\n|Abstract|中图分类号)',
                r'【关键词】\s*(.*?)(?=\n)',
            ]
        else:
            patterns = [
                r'Keywords?[:\s]*(.*?)(?=\n|JEL|\d+\.\s)',
            ]

        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                kw_text = match.group(1).strip()
                keywords = re.split(r'[;；,，\s]+', kw_text)
                return [kw.strip() for kw in keywords if kw.strip()]
        return []

    def _extract_year(self, text: str, filename: str) -> Optional[int]:
        """Extract publication year."""
        year_match = re.search(r'(20[0-2]\d|19\d{2})', filename)
        if year_match:
            return int(year_match.group(1))
        year_match = re.search(r'(20[0-2]\d|19\d{2})', text[:3000])
        if year_match:
            return int(year_match.group(1))
        return None

    def _extract_sections(self, text: str, language: str) -> list:
        """Split text into sections."""
        patterns = self.CN_SECTION_PATTERNS if language == "zh" else self.EN_SECTION_PATTERNS
        combined_pattern = '|'.join(f'({p})' for p in patterns)

        lines = text.split('\n')
        sections = []
        current_section = "Preamble"
        current_content = []

        for line in lines:
            line_stripped = line.strip()
            if line_stripped and re.match(combined_pattern, line_stripped):
                if current_content:
                    sections.append((current_section, '\n'.join(current_content)))
                current_section = line_stripped[:50]
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections.append((current_section, '\n'.join(current_content)))

        return sections


def parse_document(file_path: str) -> ParsedDocument:
    """Convenience function: parse a single document."""
    parser = PDFParser()
    return parser.parse(file_path)


def batch_parse(directory: str, pattern: str = "*.pdf") -> list:
    """Batch parse documents in a directory."""
    parser = PDFParser()
    results = []
    for pdf_path in Path(directory).glob(pattern):
        try:
            parsed = parser.parse(str(pdf_path))
            results.append(parsed)
            print(f"  [OK] {pdf_path.name}")
        except Exception as e:
            print(f"  [ERROR] {pdf_path.name} - {e}")
    return results
